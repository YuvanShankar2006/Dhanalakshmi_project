import streamlit as st
from docx import Document
from num2words import num2words
import os
def convert_number_to_words(amount_string: str) -> str:
    amount_int = int(float(amount_string))
    words = num2words(amount_int, lang='en_IN').upper()
    return words + " ONLY"
def generate_bill_docx(data, bill_filename, template_path="Dhanalakshmi traders.docx"):
    doc = Document(template_path)
    for p in doc.paragraphs:
        for key, val in data.items():
            if f"{{{{{key}}}}}" in p.text:
                for run in p.runs:
                    if f"{{{{{key}}}}}" in run.text:
                        run.text = run.text.replace(f"{{{{{key}}}}}", val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in data.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", val)
    output_path = f"{bill_filename}.docx"
    doc.save(output_path)
    return output_path
st.title("🧾 Bill Generator (.docx with amount in words)")

with st.form("bill_form"):
    bill_no = st.text_input("Bill Number")
    date = st.date_input("Date")
    from_location = st.text_input("From Location")
    to_location = st.text_input("To Location")
    vehicle_no = st.text_input("Vehicle Number")
    qty = st.number_input("Quantity", min_value=1)
    rate = st.number_input("Rate", min_value=1)
    receiver_name = st.text_input("Receiver Name")
    receiver_address = st.text_area("Receiver Address")
    receiver_state = st.text_input("Receiver State")
    receiver_gstin = st.text_input("Receiver GSTIN")

    submitted = st.form_submit_button("Generate Bill")

if submitted:
    amount = qty * rate
    igst_value = amount * 0.05
    grand_total = amount + igst_value
    date_str = date.strftime("%d/%m/%Y")
    total_in_words = convert_number_to_words(str(grand_total))

    data = {
        "receiver_name": receiver_name,
        "receiver_address": receiver_address,
        "receiver_state": receiver_state,
        "receiver_gstin": receiver_gstin,
        "bill_no": str(bill_no),
        "vehicle_no": vehicle_no,
        "date": date_str,
        "from_location": from_location,
        "to_location": to_location,
        "qty": str(qty),
        "rate": f"{rate:.2f}",
        "amount": f"{amount:.2f}",
        "taxable_amount": f"{amount:.2f}",
        "igst_value": f"{igst_value:.2f}",
        "grand_total": f"{grand_total:.2f}",
        "total_in_words": total_in_words
    }

    filename = f"bill_{bill_no}"
    filepath = generate_bill_docx(data, filename)
    st.success("✅ Bill generated successfully!")

    with open(filepath, "rb") as f:
        st.download_button(
            label="📥 Download Bill (.docx)",
            data=f,
            file_name=os.path.basename(filepath),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    os.remove(filepath)
