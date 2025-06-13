import streamlit as st
from docx import Document
from num2words import num2words
import os

def snowfall_animation():
    st.markdown("""
    <style>
    .snowflake {
        position: fixed;
        top: -10px;
        color: #fff;
        font-size: 1.5em;
        z-index: 9999;
        animation-name: fall, drift;
        animation-iteration-count: infinite;
        animation-timing-function: linear;
    }

    @keyframes fall {
        0% { top: -10px; }
        100% { top: 100vh; }
    }

    @keyframes drift {
        0% { transform: translateX(0px); }
        50% { transform: translateX(20px); }
        100% { transform: translateX(-20px); }
    }
    </style>

    <div class="snowflake" style="left: 5%; animation-duration: 10s; animation-delay: 0s;">❄</div>
    <div class="snowflake" style="left: 15%; animation-duration: 12s; animation-delay: 2s;">❅</div>
    <div class="snowflake" style="left: 25%; animation-duration: 8s; animation-delay: 4s;">❆</div>
    <div class="snowflake" style="left: 35%; animation-duration: 14s; animation-delay: 1s;">❄</div>
    <div class="snowflake" style="left: 45%; animation-duration: 11s; animation-delay: 3s;">❅</div>
    <div class="snowflake" style="left: 55%; animation-duration: 9s; animation-delay: 5s;">❆</div>
    <div class="snowflake" style="left: 65%; animation-duration: 13s; animation-delay: 2s;">❄</div>
    <div class="snowflake" style="left: 75%; animation-duration: 10s; animation-delay: 0s;">❅</div>
    <div class="snowflake" style="left: 85%; animation-duration: 15s; animation-delay: 1s;">❆</div>
    <div class="snowflake" style="left: 95%; animation-duration: 12s; animation-delay: 3s;">❄</div>
    """, unsafe_allow_html=True)


snowfall_animation()
def convert_number_to_words(amount_string: str) -> str:
    amount_int = int(float(amount_string))
    words = num2words(amount_int, lang='en_IN').upper()
    words = words.replace(",", " ").replace("-", " ")
    words = " ".join(words.split())
    return words.upper() + " ONLY"

def replace_in_paragraphs(paragraphs, data):
    for p in paragraphs:
        full_text = "".join(run.text for run in p.runs)
        replaced = full_text
        for key, val in data.items():
            replaced = replaced.replace(f"{{{{{key}}}}}", val)
        if replaced != full_text:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = replaced

def generate_bill_docx(data, bill_filename, template_path):
    doc = Document(template_path)
    replace_in_paragraphs(doc.paragraphs, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, data)
    for section in doc.sections:
        replace_in_paragraphs(section.footer.paragraphs, data)
    output_path = f"{bill_filename}.docx"
    doc.save(output_path)
    return output_path

# Party list (default)
if "party_data" not in st.session_state:
    st.session_state.party_data = {
        "SALMAN CO": "33DMZPS6425P1ZV",
        "P SHAFEEULLA AND CO.": "33ALQPS0909M1ZM",
        "S.M.D.EXPORTS": "33CGIPD6762Q1ZV",
        "SHIFA LEATHER": "33AJVPT7995F1Z7",
        "M.A.SALEEM AHMED CO.": "33AAUFM8571G1ZR",
        "BLUE WIN HIDES": "33AYIPS8536P1ZJ",
        "SARA RAJWA": "33ACEFS7879K1ZF",
        "AROON BLOSSOM IMPEX": "33ATVPS2850C1ZI",
        "M A BILAL HUSSAIN": "33AEXPH7852J1ZV",
        "BLUE TECH": "33BHAPS6391A1ZJ",
        "EVEREST EXIM": "33BDFPM7271J1ZC",
        "E.A.M EXPORTS": "33AMOPS2644E1ZZ",
        "CIPHER": "33ABEPF7243B1ZA",
        "M.S OVERSEAS": "33CTTPM4346P1ZU",
        "LOTUS INTERNATIONAL": "29AAKFL5499G1ZM",
        "H K S LEATHER": "33CCYPA5614G1ZM",
        "FIVE STAR LEATHER TEAM": "33APEPS9429Q1Z1",
        "N ABDUL WAJID CO": "33AAJFN9805F1Z8",
        "M/S NAFA ENTERPRISES": "29AHXPG6902B1Z7",
        "RAHEEM TANNING COMPANY": "33FQDPK9067B1ZW",
        "YUVAN TRADERS": "33AHIPG2293C1ZR"
    }

party_data = st.session_state.party_data
party_options = list(party_data.keys()) + ["Others"]

st.title(" Dhanalakshmi Invoice Generator")

with st.expander("Receiver Details", expanded=True):
    receiver_choice = st.selectbox("Select Receiver Name", party_options)
    if receiver_choice == "Others":
        receiver_name = st.text_input("Enter New Receiver Name")
        receiver_gstin = st.text_input("Enter New Receiver GSTIN")
        if receiver_gstin and len(receiver_gstin) != 15:
            st.warning("GSTIN must be exactly 15 characters long.")
    else:
        receiver_name = receiver_choice
        receiver_gstin = party_data[receiver_choice]
        st.markdown(f"**GSTIN:** `{receiver_gstin}`")

with st.form("bill_form"):
    tab1, tab2, tab3 = st.tabs(["Company Info", "Billing Info", "Tax and Export"])

    with tab1:
        choice = st.radio("Company", ("Yuvan Traders", "Dhanalakshmi Traders"))
        bill_no = st.text_input("Bill Number")
        date = st.date_input("Date")
        from_location = st.text_input("From Location")
        to_location = st.text_input("To Location")
        vehicle_no = st.text_input("Vehicle Number")

    with tab2:
        qty = st.number_input("Quantity", min_value=1)
        rate = st.number_input("Rate", min_value=1)
        freight = st.number_input("Freight Charges", min_value=0.0, format="%.2f")
        receiver_address = st.text_input("Receiver Address")
        receiver_state = st.text_input("Receiver State")
        Eway_bill_no = st.text_input("Eway bill Number")

    with tab3:
        is_export = st.radio("Is this an Interstate order?", ("Yes", "No"))
        dd = st.text_input("Dispatchment Details (optional)")

    submitted = st.form_submit_button("Generate Bill")

if submitted:
    if receiver_choice == "Others" and receiver_name and receiver_gstin:
        st.session_state.party_data[receiver_name] = receiver_gstin
        st.success(f"🆕 New party '{receiver_name}' added!")

    amount = qty * rate
    tamount = amount + freight
    if is_export == "Yes":
        igst_value = tamount * 0.05
        cgst_value = 0
        sgst_value = 0
        cgst = ""
        sgst = ""
        igst = "5%"
    else:
        igst_value = 0
        cgst_value = tamount * 0.025
        sgst_value = tamount * 0.025
        cgst = "2.5%"
        sgst = "2.5%"
        igst = ""

    grand_total = tamount + cgst_value + sgst_value + igst_value
    date_str = date.strftime("%d/%m/%Y")
    total_in_words = convert_number_to_words(str(grand_total))
    template = "Dhanalakshmi_new_bill.docx" if choice == "Dhanalakshmi Traders" else "YuvanTraders_bill.docx"
    data = {
        "receiver_name": receiver_name,
        "receiver_address": receiver_address,
        "receiver_state": receiver_state,
        "receiver_gstin": receiver_gstin,
        "bill_no": str(bill_no),
        "vehicle_no": vehicle_no,
        "date": date_str,
        "from_location": from_location,
        "to_location": to_location,
        "qty": str(qty),
        "rate": f"{rate:.2f}",
        "freight": f"{freight:.2f}",
        "amount": f"{amount:.2f}",
        "tamount": f"{tamount:.2f}",
        "cgst": cgst,
        "sgst": sgst,
        "igst": igst,
        "igst_value": f"{igst_value:.2f}",
        "cgst_value": f"{cgst_value:.2f}",
        "sgst_value": f"{sgst_value:.2f}",
        "grand_total": f"{grand_total:.2f}",
        "Eway_billno": Eway_bill_no,
        "total_in_words": total_in_words,
        "disp": dd
    }
    filename = f"bill_{bill_no}"
    filepath = generate_bill_docx(data, filename, template)
    st.success("✅ Bill generated successfully!")
    with open(filepath, "rb") as f:
        st.download_button(
            label="📥 Download Bill (.docx)",
            data=f,
            file_name=os.path.basename(filepath),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    os.remove(filepath)
