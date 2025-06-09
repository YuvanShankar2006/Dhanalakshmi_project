import streamlit as st
from docx import Document
import os
from tempfile import NamedTemporaryFile

def generate_bill_docx(data, template_path="Dhanalakshmi traders.docx"):
    doc = Document(template_path)
    for p in doc.paragraphs:
        for key, val in data.items():
            if f"{{{{{key}}}}}" in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if f"{{{{{key}}}}}" in inline[i].text:
                        inline[i].text = inline[i].text.replace(f"{{{{{key}}}}}", val)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in data.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", val)

    # Save to temporary .docx
    temp_docx = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_docx.name)
    return temp_docx.name

# Streamlit UI
st.title("🧾 Bill Generator (.docx version for speed)")

with st.form("bill_form"):
    bill_no = st.text_input("Bill Number", value="34")
    date = st.date_input("Date")
    from_location = st.text_input("From Location", value="KOTTAKAL")
    to_location = st.text_input("To Location", value="MELVISHARAM")
    vehicle_no = st.text_input("Vehicle Number", value="TN25Q9495")
    qty = st.number_input("Quantity", min_value=1, value=665)
    rate = st.number_input("Rate", min_value=1, value=600)
    receiver_name = st.text_input("Receiver Name", value="E.A.M EXPORTS")
    receiver_address = st.text_area("Receiver Address", value="M.V BADRAN STREET, PERIAMET CHENNAI-3")
    receiver_state = st.text_input("Receiver State", value="TAMILNADU")
    receiver_gstin = st.text_input("Receiver GSTIN", value="33AMOPS2644E1ZZ")

    submitted = st.form_submit_button("Generate Bill")

if submitted:
    amount = qty * rate
    igst_value = amount * 0.05
    grand_total = amount + igst_value
    date_str = date.strftime("%d/%m/%Y")

    data = {
        "receiver_name": receiver_name,
        "receiver_address": receiver_address,
        "receiver_state": receiver_state,
        "receiver_gstin": receiver_gstin,
        "bill_no": str(bill_no),
        "vehicle_no": vehicle_no,
        "date": date_str,
        "from_location": from_location,
        "to_location": to_location,
        "qty": str(qty),
        "rate": f"{rate:.2f}",
        "amount": f"{amount:.2f}",
        "taxable_amount": f"{amount:.2f}",
        "igst_value": f"{igst_value:.2f}",
        "grand_total": f"{grand_total:.2f}",
        "total_in_words": "AMOUNT IN WORDS HERE"
    }

    docx_path = generate_bill_docx(data)
    st.success("✅ Bill generated successfully!")

    with open(docx_path, "rb") as f:
        st.download_button(
            label="📥 Download Bill (.docx)",
            data=f,
            file_name=f"bill_{bill_no}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Clean up after the download button is rendered
    os.unlink(docx_path)
